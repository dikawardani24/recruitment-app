from __future__ import annotations

from pathlib import Path


class UnsupportedFileError(ValueError):
    pass


async def extract_text(file_name: str, content: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".pdf":
        return await _extract_pdf(content)
    if suffix in (".docx", ".doc"):
        return await _extract_docx(content)
    if suffix in (".txt", ".md", ".text"):
        return _decode(content)
    raise UnsupportedFileError(f"unsupported_file_type:{file_name}")


def _decode(content: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


async def _extract_pdf(content: bytes) -> str:
    import io

    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            chunks.append(text)
    return "\n".join(chunks)


async def _extract_docx(content: bytes) -> str:
    import io

    import docx

    document = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)
